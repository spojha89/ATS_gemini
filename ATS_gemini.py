import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import PyPDF2 as pdf
from docx import Document

load_dotenv()
genai.configure(api_key=os.getenv("gemini_API_key"))


##Function to read pdf file
def read_pdf(file):
    content=pdf.PdfReader(file)
    text=""
    for page in range(len(content.pages)):
        page=content.pages[page]
        text+=str(page.extract_text())
    return text

def read_docx(file):
    document=Document(file)
    text=""

    # Read and print each paragraph
    print("Reading paragraphs:")
    for paragraph in document.paragraphs:
        text+=str(paragraph.text)

    # Read and print each table cell
    print("\nReading tables:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text+=str(cell.text, end='\t')
            text+=str(print()) # New line after each row



##Function to get Gemini model response
def get_gemini_response(input):
    model=genai.GenerativeModel("gemini-1.5-flash")
    response=model.generate_content(input)
    return response.text

##Define prompt
input_prompt="""

Hey Act Like a skilled or very experience ATS(Application Tracking System)
with a deep understanding of tech field,software engineering,data science ,data analyst
and big data engineer. Your task is to evaluate the resume based on the given job description.
You must consider the job market is very competitive and you should provide 
best assistance for improving thr resumes. Assign the percentage Matching based 
on Jd and
the missing keywords with high accuracy
resume:{text}
description:{jd}

I want the response in one single string having the structure
{{"JD Match":"%","MissingKeywords:[]","Profile Summary":""}}

"""

##Streamlit App
st.title("ATS evaluator with Gemini")
jd=st.text_area("Mention job description")

print(jd)

uploaded_file=st.file_uploader("Select resume",type=["pdf","docx"])
submit=st.button("Submit")

if submit:
    if uploaded_file is not None:
        if uploaded_file.type=="pdf":
            text=read_pdf(uploaded_file)
        elif uploaded_file.type=="docx":
            text=read_docx(uploaded_file)
        response=get_gemini_response(input_prompt)
        st.subheader(response)
