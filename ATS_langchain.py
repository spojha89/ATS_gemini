import streamlit as st
from langchain_community.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain import LLMChain
import os
from dotenv import load_dotenv
import PyPDF2 as pdf
from docx import Document


load_dotenv()

# api_key=os.getenv("OPENAI_API_KEY")
os.environ["OPENAI_API_KEY"]='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'
# os.environ["OPENAI_API_KEY"]=os.getenv("OPENAI_API_KEY")
# print(os.getenv("OPENAI_API_KEY"))

##Function to read pdf file
def read_pdf(file):
    content=pdf.PdfReader(file)
    text=""
    for page in range(len(content.pages)):
        page=content.pages[page]
        text+=str(page.extract_text())
    return text

def read_docx(file):
    document=Document(file)
    text=""

    # Read and print each paragraph
    print("Reading paragraphs:")
    for paragraph in document.paragraphs:
        text+=str(paragraph.text)

    # Read and print each table cell
    print("\nReading tables:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text+=str(cell.text, end='\t')
            text+=str(print()) # New line after each row


##Function to get Openai model response
def get_openai_response():

    prompt=PromptTemplate(input_variables=['text','jd'],template=prompt_template)
    llm=ChatOpenAI(temperature=0.5,model_name='gpt-4o')
    llm_chain=LLMChain(llm=llm,prompt=prompt)
    response=llm_chain.run(text=text,jd=jd)

    return response

##Define prompt
prompt_template="""

Hey Act Like a skilled or very experience ATS(Application Tracking System)
with a deep understanding of tech field,software engineering,data science ,data analyst
and big data engineer. Your task is to evaluate the resume based on the given job description.
You must consider the job market is very competitive and you should provide 
best assistance for improving thr resumes. Assign the percentage Matching based 
on Jd and
the missing keywords with high accuracy
resume:{text}
description:{jd}

I want the response in one single string having the structure
{{"JD Match":"%","MissingKeywords:[]","Profile Summary":""}} 

"""

##Streamlit App
st.title("ATS evaluator with OpenAI")
jd=st.text_area("Mention job description")

uploaded_file=st.file_uploader("Select resume",type=["pdf","docx"])
submit=st.button("Submit")

text=''

if submit:
    if uploaded_file is not None:
        if uploaded_file.type=="pdf":
            text+=read_pdf(uploaded_file)
        elif uploaded_file.type=="docx":
            text+=read_docx(uploaded_file)
        print(text)
        response=get_openai_response()
        st.subheader(response)
